"""
Kubernetes Official Documentation Scraper
==========================================
Scrapes open-source Kubernetes documentation from kubernetes.io
and saves each page in PDF, DOCX, TXT, and HTML formats.

IMPORTANT — Run this locally (not in a restricted server/sandbox):
    kubernetes.io returns 403 from datacenter IPs.
    Run from your local machine or a residential proxy.

Setup:
    pip install requests beautifulsoup4 fpdf2 python-docx lxml

Usage:
    python scrape_k8s_docs.py

Output structure:
    k8s_docs/
        ├── pdf/        <- one PDF per page
        ├── docx/       <- one Word doc per page
        └── html/       <- one raw HTML file per page

Configuration (top of file):
    SEED_URLS     -- starting pages to scrape
    MAX_PAGES     -- safety cap on total pages (default: 50)
    DELAY_SECONDS -- polite delay between requests (default: 1s)
"""

import os
import re
import time
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────

SEED_URLS = [
    "https://kubernetes.io/docs/concepts/overview/what-is-kubernetes/",
    "https://kubernetes.io/docs/concepts/overview/components/",
    "https://kubernetes.io/docs/concepts/workloads/pods/",
    "https://kubernetes.io/docs/concepts/workloads/controllers/deployment/",
    "https://kubernetes.io/docs/concepts/services-networking/service/",
    "https://kubernetes.io/docs/concepts/services-networking/ingress/",
    "https://kubernetes.io/docs/concepts/storage/persistent-volumes/",
    "https://kubernetes.io/docs/concepts/scheduling-eviction/kube-scheduler/",
    "https://kubernetes.io/docs/tasks/run-application/run-stateless-application-deployment/",
    "https://kubernetes.io/docs/tasks/configure-pod-container/configure-pod-configmap/",
    "https://kubernetes.io/docs/reference/kubectl/cheatsheet/",
    "https://kubernetes.io/docs/concepts/security/pod-security-standards/",
]

OUTPUT_DIR = "data-raw/true_data/k8s_docs"
MAX_PAGES = 50
DELAY_SECONDS = 1.0

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )
}


# SETUP

def setup_dirs():
    for fmt in ["pdf", "docx", "html"]:
        os.makedirs(os.path.join(OUTPUT_DIR, fmt), exist_ok=True)
    print(f"Output directories created under '{OUTPUT_DIR}/'")


# SCRAPING UTILITIES

def slugify(url: str) -> str:
    path = urlparse(url).path
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", path).strip("_")
    return slug[:80] or "index"


def fetch_page(url: str):
    try:
        resp = requests.get(url, headers=HEADERS, timeout=15)
        resp.raise_for_status()
        return BeautifulSoup(resp.text, "lxml")
    except requests.exceptions.HTTPError as e:
        code = e.response.status_code
        if code == 403:
            print(f"  403 Forbidden — run this script locally, not in a server environment.")
        else:
            print(f"  HTTP {code} error for {url}")
        return None
    except Exception as e:
        print(f"  Failed to fetch {url}: {e}")
        return None


def extract_content(soup: BeautifulSoup) -> dict:
    title_tag = soup.find("h1") or soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else "Untitled"

    main = (
        soup.find("div", class_="td-content")
        or soup.find("main")
        or soup.find("article")
        or soup.body
    )

    sections = []
    if main:
        for tag in main.find_all(["h1", "h2", "h3", "h4", "p", "li", "pre"]):
            text = tag.get_text(separator=" ", strip=True)
            if not text or len(text) < 3:
                continue
            if tag.name in ("h1", "h2", "h3", "h4"):
                sections.append({"type": "heading", "level": int(tag.name[1]), "text": text})
            elif tag.name == "pre":
                sections.append({"type": "code", "text": text})
            else:
                sections.append({"type": "paragraph", "text": text})

    return {"title": title, "sections": sections, "url": ""}


def collect_links(soup: BeautifulSoup, current_url: str) -> list:
    seen = set()
    links = []
    for a in soup.find_all("a", href=True):
        full = urljoin(current_url, a["href"])
        parsed = urlparse(full)
        if parsed.netloc != "kubernetes.io":
            continue
        if not parsed.path.startswith("/docs/"):
            continue
        clean = parsed._replace(fragment="", query="").geturl()
        if clean not in seen:
            seen.add(clean)
            links.append(clean)
    return links


# FORMAT SAVERS

def save_html(raw_html: str, filepath: str, url: str):
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"<!-- Scraped from: {url} -->\n")
        f.write(raw_html)


def save_docx(content: dict, filepath: str):
    doc = Document()

    title_para = doc.add_heading(content["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    url_para = doc.add_paragraph(f"Source: {content['url']}")
    run = url_para.runs[0]
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(9)
    run.italic = True

    doc.add_paragraph()

    for sec in content["sections"]:
        if sec["type"] == "heading":
            doc.add_heading(sec["text"], level=min(sec["level"], 4))
        elif sec["type"] == "code":
            p = doc.add_paragraph(sec["text"])
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(9)
        else:
            doc.add_paragraph(sec["text"])

    doc.save(filepath)


PAGE_MARGIN = 15  # mm — explicit left/right margin for all pages


class KubernetesPDF(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")
        self.set_text_color(0, 0, 0)


def safe_text(text: str) -> str:
    """Encode text to latin-1 safely — FPDF2 requires this for built-in fonts."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


def get_cell_width(pdf: FPDF) -> float:
    """
    Return the usable cell width based on page width and current margins.
    Passing w=0 to multi_cell uses (page_width - left_margin - right_margin)
    but can fail if x position has drifted. This always returns a safe value.
    """
    return pdf.w - pdf.l_margin - pdf.r_margin


def save_pdf(content: dict, filepath: str):
    pdf = KubernetesPDF()

    # Set explicit margins so cell width is always predictable
    pdf.set_margins(left=PAGE_MARGIN, top=PAGE_MARGIN, right=PAGE_MARGIN)
    pdf.set_auto_page_break(auto=True, margin=PAGE_MARGIN)
    pdf.add_page()

    cell_w = get_cell_width(pdf)

    # ── Title ──
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(cell_w, 9, safe_text(content["title"]))
    pdf.ln(2)

    # ── Source URL ──
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(130, 130, 130)
    # Truncate very long URLs to prevent overflow
    url_text = safe_text(content["url"])[:120]
    pdf.multi_cell(cell_w, 5, f"Source: {url_text}")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    heading_sizes = {1: 14, 2: 12, 3: 11, 4: 10}

    for sec in content["sections"]:

        # Re-calculate cell width after every section in case x drifted
        cell_w = get_cell_width(pdf)

        if sec["type"] == "heading":
            pdf.ln(3)
            size = heading_sizes.get(sec["level"], 10)
            pdf.set_font("Helvetica", "B", size)
            pdf.multi_cell(cell_w, 7, safe_text(sec["text"]))

        elif sec["type"] == "code":
            pdf.ln(2)
            pdf.set_font("Courier", "", 7)
            pdf.set_fill_color(240, 240, 240)
            # Split long code blocks into lines to avoid single-line overflow
            code_lines = sec["text"].splitlines()
            for line in code_lines:
                # Truncate lines wider than page to prevent rendering error
                truncated = safe_text(line[:200])
                pdf.multi_cell(cell_w, 4, truncated, fill=True)
            pdf.set_fill_color(255, 255, 255)
            pdf.ln(2)

        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(cell_w, 6, safe_text(sec["text"]))
            pdf.ln(1)

    pdf.output(filepath)


# MAIN

def scrape():
    setup_dirs()

    visited = set()
    queue = list(SEED_URLS)
    saved = 0
    failed = 0

    print(f"\nStarting Kubernetes docs scrape")
    print(f"  Max pages : {MAX_PAGES}")
    print(f"  Output dir: {OUTPUT_DIR}/\n")

    while queue and saved < MAX_PAGES:
        url = queue.pop(0)

        if url in visited:
            continue
        visited.add(url)

        print(f"[{saved + 1}/{MAX_PAGES}] {url}")

        soup = fetch_page(url)
        if soup is None:
            failed += 1
            continue

        content = extract_content(soup)
        content["url"] = url

        if not content["sections"]:
            print("  No content extracted — skipping")
            continue

        slug = slugify(url)
        errors = []

        try:
            save_html(str(soup), os.path.join(OUTPUT_DIR, "html", f"{slug}.html"), url)
        except Exception as e:
            errors.append(f"HTML: {e}")

        try:
            save_docx(content, os.path.join(OUTPUT_DIR, "docx", f"{slug}.docx"))
        except Exception as e:
            errors.append(f"DOCX: {e}")

        try:
            save_pdf(content, os.path.join(OUTPUT_DIR, "pdf", f"{slug}.pdf"))
        except Exception as e:
            errors.append(f"PDF: {e}")

        if errors:
            print(f"  Save errors: {', '.join(errors)}")
        else:
            print(f"  Saved HTML / DOCX / PDF -> {slug}")

        saved += 1

        new_links = collect_links(soup, url)
        added = 0
        for link in new_links:
            if link not in visited and link not in queue:
                queue.append(link)
                added += 1
        if added:
            print(f"  {added} new links queued (queue size: {len(queue)})")

        time.sleep(DELAY_SECONDS)

    print(f"\nScrape complete")
    print(f"  Pages saved : {saved}")
    print(f"  Pages failed: {failed}")
    print(f"\n  HTML -> {OUTPUT_DIR}/html/")
    print(f"  DOCX -> {OUTPUT_DIR}/docx/")
    print(f"  PDF  -> {OUTPUT_DIR}/pdf/\n")


if __name__ == "__main__":
    scrape()